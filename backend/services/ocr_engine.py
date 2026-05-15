"""
OCR Engine — multi-format document text extraction.

Pipeline:
  Native PDFs  → PyMuPDF (text layer) + pdfplumber (table extraction → Markdown serialization)
  Scanned PDFs → OpenCV preprocessing → Tesseract quality check → GPT-4o Vision
  DOCX         → python-docx paragraph + table walk
  Images       → OpenCV preprocessing → GPT-4o Vision

Defense-in-depth OCR strategy for scanned content:
  1. Render page at 300 DPI (PyMuPDF)
  2. OpenCV preprocessing: deskew → CLAHE contrast → adaptive binarization
  3. Tesseract confidence as image quality signal — drives preprocessing intensity
  4. GPT-4o Vision OCR on the preprocessed image
  5. Confidence gate → NEEDS_MANUAL_REVIEW escalation with preprocessing metadata

Tables from government documents are serialized to GitHub-Flavoured Markdown (GFM)
so they survive chunking and are visible to the LLM extraction stage.
"""
import base64
import io
import json
import logging
import os
import re
from pathlib import Path
from typing import List, Optional, Tuple

import cv2
import fitz  # PyMuPDF
import numpy as np
import pdfplumber
import pytesseract
from docx import Document as DocxDocument
from PIL import Image
from openai import OpenAI

from backend.config import settings

logger = logging.getLogger(__name__)
client = OpenAI(api_key=settings.openai_api_key)


# ── Table serialisation ────────────────────────────────────────────────────────

def _tables_to_markdown(tables: list) -> str:
    """
    Convert pdfplumber raw table arrays (list-of-list-of-str|None) into
    GFM markdown tables so table content is preserved in the text chunks
    sent to the LLM.  Empty or degenerate tables are skipped.
    """
    if not tables:
        return ""
    md_parts = []
    for tbl in tables:
        if not tbl or not isinstance(tbl, list):
            continue
        rows = [
            [str(cell).strip() if cell is not None else "" for cell in row]
            for row in tbl
            if any(cell is not None and str(cell).strip() for cell in row)
        ]
        if len(rows) < 1:
            continue
        header = rows[0]
        col_count = len(header)
        if col_count == 0:
            continue
        lines = []
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * col_count) + " |")
        for row in rows[1:]:
            padded = (row + [""] * col_count)[:col_count]
            lines.append("| " + " | ".join(padded) + " |")
        md_parts.append("\n".join(lines))
    return "\n\n".join(md_parts)


# ── Image helpers ──────────────────────────────────────────────────────────────

def _image_to_base64(img: Image.Image) -> str:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")


# ── OpenCV Preprocessing Pipeline ─────────────────────────────────────────────

def _deskew(gray: np.ndarray) -> np.ndarray:
    """
    Deskew a grayscale image using minimum area rect on dark pixel coords.
    Skips trivial angles (< 0.5 degrees) to avoid unnecessary rotation artifacts.
    """
    coords = np.column_stack(np.where(gray < 128))
    if len(coords) < 100:
        return gray
    try:
        angle = cv2.minAreaRect(coords.astype(np.float32))[-1]
        if angle < -45:
            angle = 90 + angle
        if abs(angle) < 0.5:
            return gray
        h, w = gray.shape
        center = (w // 2, h // 2)
        M = cv2.getRotationMatrix2D(center, angle, 1.0)
        rotated = cv2.warpAffine(
            gray, M, (w, h),
            flags=cv2.INTER_CUBIC,
            borderMode=cv2.BORDER_REPLICATE,
        )
        return rotated
    except Exception:
        return gray


def _preprocess_image_for_ocr(img: Image.Image, aggressive: bool = False) -> Tuple[Image.Image, dict]:
    """
    Apply OpenCV preprocessing to improve scanned image quality before GPT-4o Vision.

    Standard pipeline:
      RGB → grayscale → deskew → CLAHE contrast enhancement → Otsu binarization

    Aggressive pipeline (triggered when Tesseract confidence is very low):
      Above steps + morphological noise removal + sharpening kernel

    Returns:
        preprocessed PIL Image (RGB), quality metadata dict
    """
    np_img = np.array(img.convert("RGB"))
    gray = cv2.cvtColor(np_img, cv2.COLOR_RGB2GRAY)

    # Step 1: Deskew
    deskewed = _deskew(gray)

    # Step 2: CLAHE contrast enhancement
    clahe = cv2.createCLAHE(clipLimit=3.0 if aggressive else 2.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(deskewed)

    if aggressive:
        # Step 3a (aggressive): morphological noise removal
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        enhanced = cv2.morphologyEx(enhanced, cv2.MORPH_CLOSE, kernel)
        # Step 3b: sharpening kernel
        sharpen_kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]], dtype=np.float32)
        enhanced = cv2.filter2D(enhanced, -1, sharpen_kernel)
        enhanced = np.clip(enhanced, 0, 255).astype(np.uint8)

    # Step 4: Otsu adaptive binarization
    _, binary = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    result_img = Image.fromarray(binary).convert("RGB")

    preprocessing_steps = ["grayscale", "deskew", "clahe"]
    if aggressive:
        preprocessing_steps += ["morphological_denoise", "sharpen"]
    preprocessing_steps.append("otsu_binarization")

    return result_img, {"preprocessing_steps": preprocessing_steps, "aggressive": aggressive}


def _get_tesseract_quality(img: Image.Image) -> float:
    """
    Run Tesseract on a PIL image and return average word-level confidence (0.0–1.0).
    Used as an early image quality signal to decide preprocessing intensity.
    Returns 0.5 as a neutral default if Tesseract fails.
    """
    try:
        np_img = np.array(img.convert("L"))  # grayscale for Tesseract
        data = pytesseract.image_to_data(
            np_img,
            output_type=pytesseract.Output.DICT,
            config="--psm 6 -c tessedit_do_invert=0",
        )
        confidences = [int(c) for c in data["conf"] if str(c) not in ("-1", "-1")]
        confidences = [c for c in confidences if c >= 0]
        if not confidences:
            return 0.5
        return round(sum(confidences) / len(confidences) / 100.0, 3)
    except Exception as e:
        logger.debug(f"Tesseract quality check failed: {e}")
        return 0.5


# ── GPT-4o Vision OCR ─────────────────────────────────────────────────────────

_OCR_TEXT_ONLY_PROMPT = (
    "You are a government-document OCR engine. "
    "Extract ALL text from this page exactly as it appears, preserving "
    "numbers, dates, clause references, and special characters. "
    "Return a JSON object with: "
    "  'text' (full extracted text, string), "
    "  'confidence' (float 0.0-1.0). "
    "Return ONLY the JSON object."
)

_OCR_TABLE_AWARE_PROMPT = (
    "You are a government-document OCR engine specialised in structured procurement forms. "
    "This page likely contains tables (financial schedules, project lists, certificate registers, "
    "EMD details, or eligibility checklists). "
    "Extract ALL content from this page. "
    "For tables, output them in GitHub-Flavoured Markdown table syntax "
    "(pipe-delimited rows with a header separator line). "
    "Preserve every cell value, row, and column header exactly. "
    "For non-table text, output it as plain text. "
    "Return a JSON object with: "
    "  'text' (all non-table prose, string), "
    "  'tables_markdown' (array of GFM table strings, one per table found on the page), "
    "  'confidence' (float 0.0-1.0 indicating overall OCR accuracy). "
    "Return ONLY the JSON object."
)


def _gpt4o_ocr_page(
    img: Image.Image,
    page_num: int,
    table_aware: bool = False,
) -> dict:
    """
    Use GPT-4o Vision to OCR a page image.

    Before calling GPT-4o:
      1. Run Tesseract to get image quality signal.
      2. Apply OpenCV preprocessing (aggressive if Tesseract confidence < 0.30).
      3. Send preprocessed image to GPT-4o Vision.

    Returns page dict with text, confidence, method, preprocessing_metadata.
    """
    # ── Quality check → preprocessing ──────────────────────────────────────
    tess_confidence = _get_tesseract_quality(img)
    aggressive = tess_confidence < 0.30
    preprocessed_img, prep_meta = _preprocess_image_for_ocr(img, aggressive=aggressive)

    image_quality_label = (
        "low" if tess_confidence < 0.30
        else "medium" if tess_confidence < 0.60
        else "high"
    )

    logger.info(
        f"Page {page_num}: Tesseract confidence={tess_confidence:.2f} "
        f"(quality={image_quality_label}), aggressive_preprocessing={aggressive}"
    )

    # ── GPT-4o Vision call ─────────────────────────────────────────────────
    b64 = _image_to_base64(preprocessed_img)
    prompt = _OCR_TABLE_AWARE_PROMPT if table_aware else _OCR_TEXT_ONLY_PROMPT
    try:
        response = client.chat.completions.create(
            model=settings.llm_model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{b64}",
                                "detail": "high",
                            },
                        },
                    ],
                }
            ],
            max_tokens=4096,
            response_format={"type": "json_object"},
        )
        raw = response.choices[0].message.content
        if not raw:
            raise ValueError("GPT-4o returned empty/null content — page may be blank or content-filtered")
        result = json.loads(raw)

        prose = result.get("text", "")
        tables_md_list = result.get("tables_markdown", [])
        tables_markdown = "\n\n".join(tables_md_list) if tables_md_list else ""

        full_text = prose
        if tables_markdown:
            full_text += "\n\n[TABLE DATA]\n" + tables_markdown

        # Weight GPT-4o self-reported confidence with Tesseract quality signal
        gpt_confidence = float(result.get("confidence", 0.7))
        # Blend: 70% GPT-4o, 30% Tesseract (Tesseract score is more objective)
        blended_confidence = round(0.7 * gpt_confidence + 0.3 * tess_confidence, 4)

        method_label = "gpt4o_vision_table_aware" if table_aware else "gpt4o_vision"
        if aggressive:
            method_label += "_aggressive_preprocess"
        elif image_quality_label != "high":
            method_label += "_preprocessed"

        return {
            "page": page_num,
            "text": full_text,
            "confidence": blended_confidence,
            "method": method_label,
            "tables_found": len(tables_md_list),
            "preprocessing_metadata": {
                "tesseract_confidence": tess_confidence,
                "image_quality": image_quality_label,
                "steps_applied": prep_meta["preprocessing_steps"],
                "aggressive_mode": aggressive,
            },
        }
    except Exception as e:
        logger.error(f"GPT-4o OCR failed on page {page_num}: {e}")
        return {
            "page": page_num,
            "text": "",
            "confidence": 0.0,
            "method": "gpt4o_vision_failed",
            "tables_found": 0,
            "preprocessing_metadata": {
                "tesseract_confidence": tess_confidence,
                "image_quality": image_quality_label,
                "steps_applied": prep_meta["preprocessing_steps"],
                "aggressive_mode": aggressive,
                "error": str(e),
            },
        }


def _is_scanned_page(page: fitz.Page, text: str) -> bool:
    """Heuristic: very little extractable text + at least one embedded image → scanned."""
    return len(text.strip()) < 50 and len(page.get_images()) > 0


def _looks_like_table_page(text: str) -> bool:
    """
    Heuristic to detect pages that are likely structured tables even when
    pdfplumber finds no formal table borders (e.g. space-aligned columns).
    """
    lines = [l for l in text.split("\n") if l.strip()]
    if not lines:
        return False
    numeric_lines = sum(
        1 for l in lines
        if re.search(r"\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\b|\b(?:crore|lakh|lac)\b", l, re.IGNORECASE)
    )
    return (numeric_lines / len(lines)) > 0.30


# ── Per-format extractors ──────────────────────────────────────────────────────

def extract_native_pdf(file_path: str) -> List[dict]:
    """
    Extract text and tables from a native (digitally created) PDF.

    Strategy per page:
      1. PyMuPDF extracts the text layer.
      2. pdfplumber extracts structured tables → serialised to GFM markdown.
      3. Table markdown is appended to the page text so it reaches the LLM.
      4. If the page is scanned (text layer empty + images present), apply
         OpenCV preprocessing + Tesseract quality check → GPT-4o Vision.
      5. If the page has no formal tables but looks like a numeric schedule,
         the table-aware GPT-4o prompt is also used.
    """
    pages = []
    doc = fitz.open(file_path)
    with pdfplumber.open(file_path) as plumber_doc:
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")
            tables = []
            try:
                plumber_page = plumber_doc.pages[page_num - 1]
                raw_tables = plumber_page.extract_tables() or []
                tables = raw_tables
            except Exception:
                pass

            if _is_scanned_page(page, text):
                # Render at 300 dpi — higher resolution helps preprocessing
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                table_aware = bool(tables) or _looks_like_table_page(text)
                ocr_result = _gpt4o_ocr_page(img, page_num, table_aware=table_aware)
                pages.append({
                    "page": page_num,
                    "text": ocr_result["text"],
                    "tables": tables,
                    "confidence": ocr_result["confidence"],
                    "method": ocr_result["method"],
                    "tables_found": ocr_result.get("tables_found", 0),
                    "preprocessing_metadata": ocr_result.get("preprocessing_metadata"),
                })
            else:
                table_md = _tables_to_markdown(tables)
                full_page_text = text
                if table_md:
                    full_page_text += "\n\n[TABLE DATA]\n" + table_md

                pages.append({
                    "page": page_num,
                    "text": full_page_text,
                    "tables": tables,
                    "confidence": 0.98,
                    "method": "native_pdf",
                    "tables_found": len(tables),
                    "preprocessing_metadata": None,
                })
    doc.close()
    return pages


def extract_docx(file_path: str) -> List[dict]:
    """
    Extract text and tables from a DOCX file.
    Tables are serialised to GFM markdown and appended to the text output.
    """
    doc = DocxDocument(file_path)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    table_mds = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            table_md = _tables_to_markdown([rows])
            if table_md:
                table_mds.append(table_md)

    combined = "\n".join(full_text)
    if table_mds:
        combined += "\n\n[TABLE DATA]\n" + "\n\n".join(table_mds)

    return [{"page": 1, "text": combined, "tables": [], "confidence": 0.99, "method": "docx_parse", "tables_found": len(table_mds), "preprocessing_metadata": None}]


def extract_image(file_path: str) -> List[dict]:
    """
    Extract text (and tables) from a standalone image via OpenCV preprocessing + GPT-4o Vision.
    """
    img = Image.open(file_path).convert("RGB")
    result = _gpt4o_ocr_page(img, 1, table_aware=True)
    return [{
        "page": 1,
        "text": result["text"],
        "tables": [],
        "confidence": result["confidence"],
        "method": result["method"],
        "tables_found": result.get("tables_found", 0),
        "preprocessing_metadata": result.get("preprocessing_metadata"),
    }]


# ── Public API ─────────────────────────────────────────────────────────────────

def extract_document(file_path: str) -> dict:
    """
    Entry point for the document extraction pipeline.

    Returns:
        pages               : list of per-page dicts (text, tables, confidence, method, preprocessing_metadata)
        full_text           : concatenated text from all pages (includes GFM table data)
        page_count          : int
        avg_confidence      : float
        primary_method      : str
        total_tables_found  : int
        low_quality_pages   : list of page numbers where image_quality was 'low' or 'medium'
    """
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".pdf":
            pages = extract_native_pdf(file_path)
        elif ext in (".docx", ".doc"):
            pages = extract_docx(file_path)
        elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"):
            pages = extract_image(file_path)
        else:
            logger.warning(f"Unsupported file type: {ext}")
            pages = [{"page": 1, "text": "", "tables": [], "confidence": 0.0, "method": "unsupported", "tables_found": 0, "preprocessing_metadata": None}]

        all_text = "\n\n".join(p["text"] for p in pages if p["text"])
        avg_confidence = sum(p["confidence"] for p in pages) / max(len(pages), 1)
        total_tables = sum(p.get("tables_found", 0) for p in pages)

        # Collect pages that needed preprocessing (low/medium quality scans)
        low_quality_pages = []
        for p in pages:
            meta = p.get("preprocessing_metadata")
            if meta and meta.get("image_quality") in ("low", "medium"):
                low_quality_pages.append({
                    "page": p["page"],
                    "image_quality": meta["image_quality"],
                    "tesseract_confidence": meta.get("tesseract_confidence"),
                    "aggressive_preprocessing": meta.get("aggressive_mode", False),
                })

        return {
            "pages": pages,
            "full_text": all_text,
            "page_count": len(pages),
            "avg_confidence": round(avg_confidence, 4),
            "primary_method": pages[0]["method"] if pages else "none",
            "total_tables_found": total_tables,
            "low_quality_pages": low_quality_pages,
        }
    except Exception as e:
        logger.error(f"Document extraction failed for {file_path}: {e}")
        return {
            "pages": [],
            "full_text": "",
            "page_count": 0,
            "avg_confidence": 0.0,
            "primary_method": "failed",
            "total_tables_found": 0,
            "low_quality_pages": [],
            "error": str(e),
        }
